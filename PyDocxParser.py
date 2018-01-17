# Copyright (C) 2009-2017 Spurgeon Woods LLC
#
# This program is free software; you can redistribute it and/or modify
# it under the terms of version 2 of the GNU General Public License as
# published by the Free Software Foundation.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program; if not, write to the Free Software
# Foundation, Inc., 59 Temple Place - Suite 330, Boston, MA 02111-1307, USA.
#

""" An XML - Docx export / import Parser for the wxPython RichTextCtrl """

__author__ = 'David Woods <dwoods@transana.com>'

DEBUG = False    # Shows debug messages

# Import wxPython and the wxPython wxRichTextCtrl
import wx
import wx.richtext as richtext

# import Python's cStringIO, os, string, and sys modules
import cStringIO, os, string, sys
# import Python's XML Sax handler
import xml.sax.handler

# import python-docx
import docx


if DEBUG:
    import time


class PyRichTextDocxHandler(richtext.RichTextFileHandler):
    """ A RichTextFileHandler that can handle Microsoft Docx files,
        at least to the extent that Transana needs that Format.
        by David K. Woods (dwoods@wcer.wisc.edu) """

    def __init__(self, name='DOCx', ext='docx'):
        """ Initialize the RichText DOCx Handler.
              Parameters:  name='DOCx'
                           ext='docx' """
        # Save the Handler Name
        self._name = name
        # Save the Handler Extension
        self._ext = ext

    def CanHandle(self, filename):
        """ Can this File Handler handle a particular file? """
        return os.path.splitext(filename)[1].lower() == ('.' + self._ext)

    def CanLoad(self):
        """ Can you load a DOCx File with this handler? """
        return True

    def CanSave(self):
        """ Can you save a DOCx File with this handler? """
        return True

    def DoLoadFile(self, buf, stream):
        return False

    def DoSaveFile(self, buf, stream):
        return False

    def GetEncoding(self):
        """ Get the encoding set for this handler """
        # NOTE:  I've only tried UTF8 encoding, which is currently hard-coded into the load and save classes.
        return 'utf8'

    def GetExtension(self):
        """ Get the handler file extension """
        return self._ext

    def GetName(self):
        """ Get the handler name """
        return self._name

    def GetType(self):
        """ Get the handler file type """
        return richtext.RICHTEXT_TYPE_ANY

    def IsVisible(self):
        return True

    def LoadFile(self, ctrl, filename):
        """ Load the contents of a Docx Format file into a wxRichTextCtrl.
            Parameters:  ctrl       a wxRichTextCtrl.  (NOT a wxRichTextBuffer.  The wxRichTextBuffer lacks methods for direct manipulation.)
                         filename   the name of the file to be loaded """
        if os.path.exists(filename) and isinstance(ctrl, richtext.RichTextCtrl):
            # Use the DocxToRichTextCtrlParser to handle the file load
            DocxTowxRichTextCtrlParser(ctrl, filename=filename, encoding=self.GetEncoding())
            # There's no feedback from the Parser, so we'll just assume things loaded.
            return True
        else:
            return False

##    def LoadString(self, ctrl, buf, insertionPoint=None, displayProgress=True):
##        """ Load the contents of a Rich Text Format string buffer into a wxRichTextCtrl.
##            Parameters:  ctrl       a wxRichTextCtrl.  (NOT a wxRichTextBuffer.  The wxRichTextBuffer lacks methods for direct manipulation.)
##                         buf        the DOCx string data to be loaded """
##        if (len(buf) > 0) and isinstance(ctrl, richtext.RichTextCtrl):
##
##            # At least in Transana, if the buffer is a unicode object, processing is MUCH, MUCH slower, like more than
##            # 20 TIMES slower, than if the object is a string.  Converting from unicode to string speeds things up incredibly.
##            # At least for Transana, this causes no problems.
##            if isinstance(buf, unicode):
##                buf = buf.encode(self.GetEncoding())
##                
##            # Use the DocxToRichTextCtrlParser to handle the file load
##            DocxTowxRichTextCtrlParser(ctrl, buf=buf, insertionPoint=insertionPoint, encoding=self.GetEncoding(), displayProgress=displayProgress)
##            # There's no feedback from the Parser, so we'll just assume things loaded.
##            return True
##        else:
##            return False


    def SaveFile(self, buf, filename=None):
        """ Save the contents of a wxRichTextBuffer to a Rich Text Format file.
            Parameters:  buf       a wxRichTextBuffer or a wxRichTextCtrl
                         filename  the name of the file to be created or overwritten """
        # If we're passed a wxRichTextCtrl, we can get the control's buffer, which is what we need.
        if isinstance(buf, richtext.RichTextCtrl):
            buf = buf.GetBuffer()

        # Get a Rich Text XML Handler to extract the data from the wxRichTextBuffer in XML.
        # NOTE:  buf.Dump() just returns the text contents of the buffer, not any formatting information.
        xmlHandler = richtext.RichTextXMLHandler()
        # Create a stream object that can hold the data
        stream = cStringIO.StringIO()

        # If no file name is specified, we should return a string instead of saving to a file!
        if filename == None:

            print "ERROR -- Filename is REQUIRED for docx export!"
            
            return
        # If a filename is specified ..
        else:
            # ... we can just pass the file name to the file handler
            fileobj = filename

        # Extract the wxRichTextBuffer data to the stream object
        if xmlHandler.SaveStream(buf, stream):
            # Convert the stream to a string
            contents = stream.getvalue()
            # Get the XML to Docx File Handler
            fileHandler = XMLToDocxHandler()
            # Use xml.sax, with the XML to Docx File Handler, to parse the XML and create
            # a Docx Output string.
            xml.sax.parseString(contents, fileHandler)
            # Use the XML to Docx File Handler to save the Docx Output String to a file or
            # to populate the StringIO object if we're to return a string
            fileHandler.saveFile(fileobj)

            # If no file name is specified ...
            if filename == None:
                # ... return the converted Docx String
                return fileobj.getvalue()
            # If a filename is specified ...
            else:
                # ... indicate success in saving
                return True
        # If we couldn't extract the XML from the buffer ...
        else:
            # ... signal failure
            return False

    def SetName(self, name):
        """ Set the name of the File Handler """
        self._name = name



class XMLToDocxHandler(xml.sax.handler.ContentHandler):
    """ An xml.sax handler designed to convert wxRichTextCtrl's internal XML format data into
        Word docx Format data that can be saved to *.docx files, at least to the extent that
        Transana (htp://www.transana.org) needs Word docx Format features supported.
        by David K. Woods (dwoods@wcer.wisc.edu) """

    def __init__(self, encoding='utf8'):
        """ Initialize the XMLToDocxHandler
            Parameters:  encoding='utf8'  Character Encoding to use (only utf8 has been tested, and I don't
                                          think the Docx Parser decodes yet. """
        # Remember the encoding to use
        self.encoding = encoding

        # Define an initial Fonts.  We define multiple levels of fonts to handle cascading styles.
        self.fontAttributes = {}
        self.fontAttributes[u'text'] = {u'bgcolor' : '#FFFFFF',
                                    u'fontface' : 'Courier New',
                                    u'fontpointsize' : 11,  # Word defaults to 11 pt, I believe
                                    u'fontstyle' : wx.FONTSTYLE_NORMAL,
                                    u'fontunderlined' : u'0',
                                    u'fontweight' : wx.FONTSTYLE_NORMAL,
                                    u'textcolor' : '#000000'}

        self.fontAttributes[u'symbol'] = {u'bgcolor' : '#FFFFFF',
                                    u'fontface' : 'Courier New',
                                    u'fontpointsize' : 11,  # Word defaults to 11 pt, I believe
                                    u'fontstyle' : wx.FONTSTYLE_NORMAL,
                                    u'fontunderlined' : u'0',
                                    u'fontweight' : wx.FONTSTYLE_NORMAL,
                                    u'textcolor' : '#000000'}

        self.fontAttributes[u'paragraph'] = {u'bgcolor' : '#FFFFFF',
                                         u'fontface' : 'Courier New',
                                         u'fontpointsize' : 11,  # Word defaults to 11 pt, I believe
                                         u'fontstyle' : wx.FONTSTYLE_NORMAL,
                                         u'fontunderlined' : u'0',
                                         u'fontweight' : wx.FONTSTYLE_NORMAL,
                                         u'textcolor' : '#000000'}

        self.fontAttributes[u'paragraphlayout'] = {u'bgcolor' : '#FFFFFF',
                                               u'fontface' : 'Courier New',
                                               u'fontpointsize' : 11,  # Word defaults to 11 pt, I believe
                                               u'fontstyle' : wx.FONTSTYLE_NORMAL,
                                               u'fontunderlined' : u'0',
                                               u'fontweight' : wx.FONTSTYLE_NORMAL,
                                               u'textcolor' : '#000000'}

        # Define the initial Paragraph attributes.  We define mulitple levels to handle cascading styles.
        self.paragraphAttributes = {}
        self.paragraphAttributes[u'paragraph'] = {u'alignment' : u'1',
                                                  u'linespacing' : u'10',
                                                  u'leftindent' : u'0',
                                                  u'rightindent' : u'0',
                                                  u'leftsubindent' : u'0',
                                                  u'parspacingbefore' : u'0',
                                                  u'parspacingafter' : u'35',  # Word defaults to 10pt, which is 35!
                                                  u'bulletnumber' : None,
                                                  u'bulletstyle' : None,
                                                  u'bulletfont' : None,
                                                  u'bulletsymbol' : None,
                                                  u'bullettext' : None,
                                                  u'tabs' : None}

        self.paragraphAttributes[u'paragraphlayout'] = {u'alignment' : u'1',
                                                      u'linespacing' : u'10',
                                                      u'leftindent' : u'0',
                                                      u'rightindent' : u'0',
                                                      u'leftsubindent' : u'0',
                                                      u'parspacingbefore' : u'0',
                                                      u'parspacingafter' : u'35',  # Word defaults to 10pt, which is 35!
                                                      u'bulletnumber' : None,
                                                      u'bulletstyle' : None,
                                                      u'bulletfont' : None,
                                                      u'bulletsymbol' : None,
                                                      u'bullettext' : None,
                                                      u'tabs' : None}

        # Getting this to work both from within Python and in the stand-alone executable
        # has been a little tricky.  To get it working right, we need the path to the
        # Transana executables, where the python-docx default template can be found.

        # If we're working from a stand-alone executable ...
        if hasattr(sys, "frozen"):
            # Let's find the path
            (path, fn) = os.path.split(sys.argv[0])
            # If the path is not blank, add the path seperator to the end if needed
            if (path != '') and (path[-1] != os.sep):
                path = path + os.sep

            # Create a python-docx Document using the default document template.  This MUST be defined when using py2exe
            self.document = docx.Document(path + 'images/default.docx')
        # If running from source code ...
        else:
            # ... then docx knows where to find it's default template automatically.
            self.document = docx.Document()

        # Define a variable for tracking what element we are changing
        self.element = ''

        # Define a variable for tracking images
        self.image_num = 0

        # Highlight Colors defined in Microsoft Word and in Transana need to be translated into python_docx constants.
        # I find a few of these translations puzzling, but this is what Word does given fully-defined colors in DOCs files.
        enum = docx.enum.text.WD_COLOR_INDEX
        self.colorConversions = { # enum.AUTO          :  None,                     # Word:  Automatic
                                 "#000000": enum.BLACK,          # Word:  Black
                                 "#0000FF": enum.BLUE,           # Word:  Blue,        Transana: Blue
                                 "#006400": enum.GREEN,          # Word:  Green
                                 "#0080FF": enum.TURQUOISE,      #                     Transana: Light Blue
                                 "#008000": enum.GREEN,          #                     Transana: Dark Green
                                 "#008080": enum.TEAL,           #                     Transana: Blue Green
                                 "#008B8B": enum.TEAL,           # Word:  Teal
                                 "#00FF00": enum.BRIGHT_GREEN,   # Word:  Bright Green
                                 "#000080": enum.DARK_BLUE,      #                     Transana: Dark Blue
                                 "#00008B": enum.DARK_BLUE,      # Word:  Dark Blue
                                 "#00FF80": enum.BRIGHT_GREEN,   #                     Transana: Green Blue
                                 "#00FFFF": enum.TURQUOISE,      # Word:  Turquoise,   Transana: Cyan
                                 "#2F4F4F": enum.BLACK,          #                     Transana: Dark Slate Gray
                                 "#4F2F2F": enum.DARK_RED,       #                     Transana: Indian Red
                                 "#800000": enum.DARK_RED,       #                     Transana: Maroon
                                 "#800080": enum.VIOLET,         # Word:  Violet       Transana: Dark Purple
                                 "#8000FF": enum.PINK,           #                     Transana: Purple
                                 "#808000": enum.DARK_YELLOW,    # Word:  Dark Yellow  Transana: Olive
                                 "#808080": enum.GRAY_50,        #                     Transana: Gray
                                 "#8080FF": enum.BLUE,           #                     Transana: Lavendar
                                 "#80FF00": enum.YELLOW,         #                     Transana: Chartruese
                                 "#80FF80": enum.BRIGHT_GREEN,   #                     Transana: Light Green
                                 "#80FFFF": enum.TURQUOISE,      #                     Transana: Light Aqua
                                 "#8B0000": enum.DARK_RED,       # Word:  Dark Red
                                 "#8E6B23": enum.DARK_YELLOW,    #                     Transana: Sienna
                                 "#A9A9A9": enum.GRAY_50,        # Word:  Gray 50%
                                 "#B000FF": enum.PINK,           #                     Transana: Light Purple
                                 "#C0C0C0": enum.GRAY_25,        #                     Transana: Light Gray
                                 "#CC3232": enum.RED,            #                     Transana: Red Orange
                                 "#CC3299": enum.PINK,           #                     Transana: Violet Red
                                 "#D3D3D3": enum.GRAY_25,        # Word:  Gray 25%
                                 "#DBDB70": enum.YELLOW,         #                     Transana: Goldenrod
                                 "#FF0000": enum.RED,            # Word:  Red          Transana: Red
                                 "#FF0080": enum.PINK,           #                     Transana: Rose
                                 "#FF00FF": enum.PINK,           # Word:  Pink         Transana: Magenta
                                 "#FF8000": enum.YELLOW,         #                     Transana: Orange
                                 "#FF8080": enum.RED,            #                     Transana: Salmon
                                 "#FF80FF": enum.PINK,           #                     Transana: Light Fuchsia
                                 "#FFFF00": enum.YELLOW,         # Word:  Yellow       Transana: Yellow
                                 "#FFFF80": enum.YELLOW,         #                     Transana: Light Yellow
                                 "#FFFFFF": enum.WHITE}          # Word:  White


    def startElement(self, name, attributes):
        """ xml.sax required method for handling the starting XML element """

##        # We need roman numerals for list processing
##        # Copied from http://www.daniweb.com/code/snippet216865.html on 2/3/2010
##        def int2roman(number):
##            numerals = { 1 : "I", 4 : "IV", 5 : "V", 9 : "IX", 10 : "X", 40 : "XL",
##                        50 : "L", 90 : "XC", 100 : "C", 400 : "CD", 500 : "D", 900 : "CM", 1000 : "M" }
##            result = ""
##            for value, numeral in sorted(numerals.items(), reverse=True):
##                while number >= value:
##                    result += numeral
##                    number -= value
##            return result

        # Remember the element's name
        self.element = name

        if DEBUG:
            print "PyDocxParser.XMLToDocxHandler.startElement():", name

        # If the element is a paragraphlayout, paragraph, symbol, or text element ...
        if name in [u'paragraphlayout', u'paragraph', u'symbol', u'text']:

            # Let's cascade the font and paragraph settings from a level up BEFORE we change things to reset the font and  
            # paragraph settings to the proper initial state.  First, let's create empty character and paragraph cascade lists
            charcascade = paracascade = []
            # Initially, assume we will cascade from our current object for character styles
            cascadesource = name
            
            # If we're in a Paragraph spec ...
            if name == u'paragraph':
                # ... we need to cascase paragraph, symbol, and text styles for characters ...
                charcascade = [u'paragraph', u'symbol', u'text']
                # ... from the paragraph layout style for characters ...
                cascadesource = u'paragraphlayout'
                # ... and we need to cascare paragraph styles for paragraphs
                paracascade = [u'paragraph']
            # If we're in a Text spec ...
            elif name == u'text':
                # ... we need to cascase text styles for characters ...
                charcascade = [u'text']
                # ... from the paragraph style for characters ...
                cascadesource = u'paragraph'
            # If we're in a Symbol spec ...
            elif name == u'symbol':
                # ... we need to cascase symbol styles for characters ...
                charcascade = [u'symbol']
                # ... from the paragraph style for characters ...
                cascadesource = u'paragraph'
            # For each type of character style we need to cascade ...
            for x in charcascade:
                # ... iterate through the dictionary elements ...
                for y in self.fontAttributes[x].keys():
                    # ... and assign the character cascade source styles (cascadesource) to the destination element (x)
                    self.fontAttributes[x][y] = self.fontAttributes[cascadesource][y]
            # For each type of paragraph style we need to cascade ...
            for x in paracascade:
                # ... iterate through the dictionary elements ...
                for y in self.paragraphAttributes[x].keys():
                    # ... and assign the paragraph cascade source styles (cascadesource) to the destination element (x)
                    self.paragraphAttributes[x][y] = self.paragraphAttributes[cascadesource][y]

            # If the element is a paragraph element or a paragraph layout element, there is extra processing to do at the start
            if name in [u'paragraph', u'paragraphlayout']:
                # ... iterate through the element attributes looking for paragraph attributes
                for x in attributes.keys():
                    # If the attribute is a paragraph format attribute ...
                    if x in [u'alignment',
                             u'linespacing',
                             u'leftindent',
                             u'rightindent',
                             u'leftsubindent',
                             u'parspacingbefore',
                             u'parspacingafter',
                             u'bulletnumber',
                             u'bulletstyle',
                             u'bulletfont',
                             u'bulletsymbol',
                             u'bullettext',
                             u'tabs']:
                        # ... update the current paragraph dictionary
                        self.paragraphAttributes[name][x] = attributes[x]

            # ... iterate through the element attributes looking for font attributes
            for x in attributes.keys():
                if x == u'fontsize':
                    x = u'fontpointsize'
                # If the attribute is a font format attribute ...
                if x in [u'bgcolor',
                         u'fontface',
                         u'fontpointsize',
                         u'fontstyle',
                         u'fontunderlined',
                         u'fontweight',
                         u'textcolor']:
                    # ... update the current font dictionary
                    self.fontAttributes[name][x] = attributes[x]

##                # If the element is a text element and the attribute is a url attribute ...
##                if (name == u'text') and (x == u'url'):
##                    # ... capture the URL data.
##                    self.url = attributes[x]
##
##            # If the URL is a Transana Object link ...
##            # (This should be done after all text attributes are processed so formatting can be corrected.)
##            if (len(self.url) > 9) and (self.url[:9].lower() == 'transana:'):
##                # ... completely remove the URL value
##                self.url = ''
##                # Let's remove the Hyperlink formatting too!
##                self.fontAttributes[u'text'][u'textcolor'] = '#000000'
##                self.fontAttributes[u'text'][u'fontunderlined'] = u'0'

            # Let's cascade the font and paragraph settings we've just changed.
            # First, let's create empty character and paragraph cascade lists
            charcascade = paracascade = []
            # Initially, assume we will cascade from our current object for character styles
            cascadesource = name
            # If we're in a Paragraph Layout spec ...
            if name == u'paragraphlayout':
                # ... we need to cascase paragraph, symbol, and text styles for characters ...
                charcascade = [u'paragraph', u'symbol', u'text']
                # ... we need to cascase paragraph styles for paragraphs ...
                paracascade = [u'paragraph']
            # If we're in a Paragraph spec ...
            elif name == u'paragraph':
                # ... we need to cascase symbol and text styles for characters ...
                charcascade = [u'symbol', u'text']
            # For each type of character style we need to cascade ...
            for x in charcascade:
                # ... iterate through the dictionary elements ...
                for y in self.fontAttributes[x].keys():
                    # ... and assign the character cascade source styles (cascadesource) to the destination element (x)
                    self.fontAttributes[x][y] = self.fontAttributes[cascadesource][y]
            for x in paracascade:
                # ... iterate through the dictionary elements ...
                for y in self.paragraphAttributes[x].keys():
                    # ... and assign the paragraph cascade source styles (cascadesource) to the destination element (x)
                    self.paragraphAttributes[x][y] = self.paragraphAttributes[cascadesource][y]

        # If the element is an image element ...
        elif name in [u'image']:
            # ... if we have a PNG graphic ...
            if attributes[u'imagetype'] == u'15':  # wx.BITMAP_TYPE_PNG = 15
                # ... signal that we have a PNG image to process ...
                self.elementType = "ImagePNG"
            # It appears to me that all images will be PNG images coming from the RichTextCtrl.
            else:
                # if not, signal a unknown image type
                self.elementType = 'ImageUnknown'

                print "Image of UNKNOWN TYPE!!", attributes.keys()

        # If the element is a data element ...
        elif name in [u'data']:
            # ... we *probably* have image data.  Sometimes image data is too large to fit
            #     into a single data element and gets split across several.  Therefore, we'll
            #     initialize a variable to collect that data here and process it in the
            #     endElement method.
            self.dataValue = ''

        # If the element is a richtext element ...
        elif name in [u'richtext']:
            # ... we should not have to do anything!
            pass
        
        # If we have an unhandled element ...
        else:
            # ... output a message and the element attributes.
            print "PyDocxParser.XMLToDocxHandler.startElement():  Unknown XML tag:", name
            for x in attributes.keys():
                print x, attributes[x]
            print

        # If the element is a paragraph element ...
        if name in [u'paragraph']:


            # Create a docx Paragraph
            self.paragraph = self.document.add_paragraph()

            # Assign the appropriate paragraph formatting
            par_format = self.paragraph.paragraph_format
            
            # Paragraph alignment left is u'1'
            if self.paragraphAttributes[u'paragraph'][u'alignment'] == u'1':
                par_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            # Paragraph alignment centered is u'2'
            elif self.paragraphAttributes[u'paragraph'][u'alignment'] == u'2':
                par_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            # Paragraph alignment right is u'3'
            elif self.paragraphAttributes[u'paragraph'][u'alignment'] == u'3':
                par_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            else:
                print "Unknown alignment:", self.paragraphAttributes[u'paragraph'][u'alignment'], type(self.paragraphAttributes[u'paragraph'][u'alignment'])

            # line spacing
            if self.paragraphAttributes[u'paragraph'][u'linespacing'] in [u'0']:
                pass
            else:
                # Line Spacing is one tenth of the value here, with 10 = single spacing, 20 = double spacing, etc.
                par_format.line_spacing = float(self.paragraphAttributes[u'paragraph'][u'linespacing']) / 10.0

            # Paragraph Margins and first-line indents
            # First, let's convert the unicode strings we got from the XML to integers and translate from wxRichTextCtrl's
            # system to Docs's system.
            # Left Indent in Docx is the sum of wxRichTextCtrl's left indent and left subindent
            leftindent = int(self.paragraphAttributes[u'paragraph'][u'leftindent']) + int(self.paragraphAttributes[u'paragraph'][u'leftsubindent'])
            # The First Line Indent in Docx is the wxRichTextCtrl's left indent minus the left indent calculated above.
            firstlineindent = int(self.paragraphAttributes[u'paragraph'][u'leftindent']) - leftindent
            # The Right Indent translates directly
            rightindent = int(self.paragraphAttributes[u'paragraph'][u'rightindent'])

            # All values are in 1/10ths of a millimeter.
            # Now add these values to the Docx paragraph format
            par_format.left_indent = docx.shared.Mm(leftindent / 10.0)
            par_format.first_line_indent = docx.shared.Mm(firstlineindent / 10.0)
            par_format.right_indent = docx.shared.Mm(rightindent / 10.0)

            # Add non-zero Spacing before and after paragraphs to the DOCx output String
            if int(self.paragraphAttributes[u'paragraph'][u'parspacingbefore']) != 0:
                par_format.space_before = docx.shared.Mm(float(self.paragraphAttributes[u'paragraph'][u'parspacingbefore']) / 10.0)
            if int(self.paragraphAttributes[u'paragraph'][u'parspacingafter']) > 0:
                par_format.space_after = docx.shared.Mm(float(self.paragraphAttributes[u'paragraph'][u'parspacingafter']) / 10.0)
            # Due to a bug in the RichTextEditCtrl, the parspacingafter value may sometimes be NEGATIVE, which of course doesn't
            # make sense outside of the RichTextEditCtrl.  This adjusts for that.
            else:
                parAfter = int(self.paragraphAttributes[u'paragraph'][u'parspacingafter']) # + int(self.paragraphAttributes[u'paragraph'][u'parspacingbefore'])
                par_format.space_after = docx.shared.Mm(max(parAfter, 0) / 10.0)
            # If Tabs are defined ...
            if self.paragraphAttributes[u'paragraph'][u'tabs'] != None:
                # ... break the tab data into its component pieces
                tabStops = self.paragraphAttributes[u'paragraph'][u'tabs'].split(',')
                # For each tab stop ...
                for x in tabStops:
                    # ... (assuming the data isn't empty) ...
                    if x != u'':
                        # ... add the tab stop data to the DOCx output
                        par_format.tab_stops.add_tab_stop(docx.shared.Mm(int(x) / 10.0))

    def characters(self, data):
        """ xml.sax required method for handling the characters within XML elements """

        # If we have a text or symbol element ...
        if self.element in ['text', 'symbol']:

            # If we have whitespace only ...
            if len(data.strip()) == 0:  # len(data) == 1:
                pass

            # If we have data ...
            else:
                # If the text has leading or trailing spaces, it gets enclosed in quotation marks in the XML.
                # Otherwise, not.  We have to detect this and remove the quotes as needed.  Unicode characters
                # make this a bit more complicated, as in " 137 e(umlaut) 137 ".
                if ((data != ' "') and ((data[0] == '"') or (data[-1] == '"'))  ):
                    if data[0] == '"':
                        data = data[1:]
                    if (len(data) > 1) and (data[-1] == '"'):
                        data = data[:-1]


                # Create a run
                run = self.paragraph.add_run()

                # Assign the appropriate character formatting
                run.font.name = self.fontAttributes[self.element][u'fontface']
                # Add Font Size information
                run.font.size = docx.shared.Pt(int(self.fontAttributes[self.element][u'fontpointsize']))
                
                # If bold, add Bold
                if self.fontAttributes[self.element][u'fontweight'] == str(wx.FONTWEIGHT_BOLD):
                    run.bold = True
                # If Italics, add Italics
                if self.fontAttributes[self.element][u'fontstyle'] == str(wx.FONTSTYLE_ITALIC):
                    run.italic = True
                # If Underline, add Underline
                if self.fontAttributes[self.element][u'fontunderlined'] == u'1':
                    run.underline = True
                # If Text Color is not black ...
                if self.fontAttributes[self.element][u'textcolor'] != '#000000':
                    # ... Add text foreground color
                    run.font.color.rgb = docx.shared.RGBColor.from_string("%s" % self.fontAttributes[self.element][u'textcolor'][1:])

                # If Text Background Color is defined in the docx Word Highlight Colors enumeration ...
                if (self.fontAttributes[self.element][u'bgcolor'].upper() in self.colorConversions.keys()):
                    # If Text Background Color is not White ...
                    if (self.fontAttributes[self.element][u'bgcolor'] != '#FFFFFF'):
                        # ... Add text background color to the DOCx output string
                        run.font.highlight_color = self.colorConversions[self.fontAttributes[self.element][u'bgcolor'].upper()]

                else:
                    print "UNDEFINED BACKGROUND COLOR!!", self.fontAttributes[self.element][u'bgcolor']
        
        # If the characters come from a text element ...
        if self.element in ['text']:

##            # If we have a value in self.URL, populated in startElement, ...
##            if self.url != '':
##                # ... then we're in the midst of a hyperlink.  Let's specify the URL for the DOCx output string.
##                self.outputString.write('{\\field{\\*\\fldinst HYPERLINK "%s"}{\\fldrslt ' % self.url)

            if DEBUG:
                print "  -->  ", data.encode('utf8')


            # If we have more than whitespace ...
            if len(data.strip()) > 0:  # len(data) == 1:
                # ... add it as the run text
                run.text = data
                    
##                # If we're in Transana, time code data if followed by a "(space)(quotationmark)" combination from the XML.
##                # I'm not sure why, but this causes problems in the DOCx.  Therefore skip this combo in Transana
##                if not (IN_TRANSANA and (data == ' "')):
##                    # Encode the data and add it to the DOCx output string
##                self.outputString.write(data.encode(self.encoding))

##            # If we've just added a URL hyperlink ...
##            if self.url != '':
##                # ... we need to close the link field DOCx block
##                self.outputString.write('}}')
##                # Reset the URL to empty, as we're done with it.
##                self.url = ''

        # If the characters come from a symbol element ...
        elif self.element == 'symbol':

            if DEBUG:
                print "PyDocxParser.characters(): element == symbol: '%s'" % data
            
            # Check that we don't have only whitespace, we don't have a multi-character string, and
            # we don't have a newline character.
            if (len(data.strip()) > 0) and ((len(data) != 1) or (ord(data) != 10)):

                try:
                    # Convert the symbol data to the appropriate unicode character
                    data = unichr(int(data))
                    # Add that unicode character to the DOCx output string
                    run.text = data
                except ValueError, e:

                    if DEBUG:
                        print "PyDocxParser.XMLToDocxHandler.characters():  ValueError", len(data)
                    
                        for x in range(len(data)):
                            print x, ord(data[x])
                        print
                        
                    if len(data) == 1:
                        try:
                            run.text = unichr(ord(data))
                        except:
                            run.text = '?'
                    else:
                        run.text = '(Unknown Character)'
                except:
                    run.text = '(Unknown Character)'

                    print "Unknown Character: '%s' - %d" % (data, int(data))
                    print


        # If the characters come from a data element ...
        elif self.element == 'data':
            # ... we *probably* have image data.  Sometimes image data is too large to fit
            #     into a single data element and gets split across several.  Therefore, we'll
            #     collect that data here and process it in the endElement method.
            self.dataValue += data

        # We can ignore whitespace here, which will be made up of the spaces added by XML and newline characters
        # that are part of the XML file but not part of the data.
        elif data.strip() != '':
            # Otherwise, print a message to the developer
            print "PyDocxParser.characters():  Unhandled text."
            print '"%s"' % data
            
    def endElement(self, name):
        """ xml.sax required method for handling the ending of an XML element (the close tag) """

        # If we have a text, paragraph, paragraphlayout, or richtext end tag ...
        if name in [u'text', u'paragraph', u'paragraphlayout', u'richtext']:
            # ... we need to clear the element type, as we're no longer processing that type of element!
            self.element = None

        # If we have a data element ...
        elif name in [u'data']:
            # ... we *probably* have image data.  Sometimes image data is too large to fit
            #     into a single data element and gets split across several.  Therefore, we'll
            #     collect that data abpve and process it here in the endElement method.

            # If we're expecting a PNG Image ...
            if self.elementType == 'ImagePNG':

                # Create a StringIO stream from the HEX-converted image data (which is a PNG image)
                stream = cStringIO.StringIO(self.hex2int(self.dataValue))
                # Create a run
                run = self.paragraph.add_run()
                # Add the image to the run
                run.add_picture(stream)
                # Close the stream
                stream.close()
               
            # I haven't seen anything but PNG image data in this data structure from the RichTextCtrl's XML data
            else:
                # If we're dealing with an image, we could convert the image to PNG, then do a Hex conversion.
                # DOCx can also JPEG images directly, as well as Enhanced Metafiles, Windows Metafiles, QuickDraw
                # pictures, none of which I think wxPython can handle.
                print "I don't know how to handle the data!!"

            # ... we need to clear the element type, as we're no longer processing that type of element!
            self.element = None

    def saveFile(self, filename):
        """ Save the DOCx Output String to a file or to a StringIO object """

        try:
            self.document.save(filename)
        except:

            msg = "PyDocxParser.XMLtoDocxHandler.saveFile():  Could not save file....\n\n%s\n\n%s" % (sys.exc_info()[0], sys.exc_info()[1])
            import Dialogs
            dlg = Dialogs.InfoDialog(None, msg)
            dlg.ShowModal()
            dlg.Destroy()

            print "Could not save DOCX file."
            print sys.exc_info()[0]
            print sys.exc_info()[1]
            print
        
    def hex2int(self, data):
        """ Image data is stored in a file-friendly Hex format.  We need to convert it to an image-friendly binary format. """
        # Initialize the conversion result variable
        result = ''
        # For each PAIR of characters in the hex data string ...
        for x in range(0, len(data), 2):
            # ... convert the hex pair into a integer, find that character, and add it to the result variable
            result += chr(int(data[x : x + 2], 16))
        # Return the converted data
        return result


## For Docx Import

class DocxTowxRichTextCtrlParser:
    """ A Docx Parser designed to convert Microsoft Docx data from *.docx files to
        wxRichTextCtrl's internal format, at least to the extent that
        Transana (htp://www.transana.org) needs Docx features supported.
        by David K. Woods (dwoods@wcer.wisc.edu) """

    def __init__(self, txtCtrl, filename=None, buf=None, insertionPoint=None, encoding='utf8', displayProgress=True):
        """ Initialize the DocxToRichTextCtrlParser.

            Parameters:  txtCtrl          a wx.RichTextCtrl, NOT a wx.RichTextBuffer.  The buffer doesn't provide an easy way to add text!
                         filename=None    a Microsoft Docx Format (*.docx) file name
                         XXXX buf=None         a string with Docx-encoded data
                         XXXX insertionPoint   NOT IMPLEMENTED
                         encoding='utf8'  Character Encoding to use (only utf8 has been tested, and I don't
                                          think the Docx Parser decodes yet.

            You can pass in either a filename or a buffer string.  If both are passed, only the file will be imported.  """

        # Remember the wxRichTextCtrl to populate
        self.txtCtrl = txtCtrl
        # Remember the insertion point
        self.insertionPoint = insertionPoint
        if insertionPoint != None:
            self.insertionOffset = self.txtCtrl.GetLastPosition() - insertionPoint
        else:
            self.insertionOffset = 0

        # At present, encoding is not used!
        self.encoding = encoding

        # Create a default font specification.  I've chosen Courier New, 12 point, black on white,
        self.font = {'fontfacename'  :  'Courier New',
                     'fontpointsize' :  11,  # Word defaults to 11 pt, I believe
                     'fontcolor'     :  wx.Colour(0, 0, 0),
                     'fontbgcolor'   :  wx.Colour(255, 255, 255)}

        # Initialize Paragraph settings
        self.paragraph = {'alignment'       : 'left',
                          'linespacing'     : wx.TEXT_ATTR_LINE_SPACING_NORMAL,
                          'leftindent'      : 0,
                          'rightindent'     : 0,
                          'firstlineindent' : 0,
                          'spacingbefore'   : 0,
                          'spacingafter'    : 35,  # Word defaults to 10pt, which is 35!
                          'tabs'            : []}

        # Create an object to hold font specifications for the current font
        self.txtAttr = richtext.RichTextAttr()
        
        # Apply the default font specifications to the current font object
        self.SetTxtStyle(fontFace = self.font['fontfacename'], fontSize = self.font['fontpointsize'],
                          fontColor = self.font['fontcolor'], fontBgColor = self.font['fontbgcolor'],
                          fontBold = False, fontItalic = False, fontUnderline = False)

        # If a file name was passed in and the file exists ...
        if (filename != None) and os.path.exists(filename):
            # ... open the file to be read ...
            self.document = docx.Document(filename)

        # If there's nothing to read ...
        else:
            # ... create an empty Document
            self.document = docx.Document()

        # Set the insertion point, if one is passed in
        if self.insertionPoint != None:
            self.txtCtrl.SetInsertionPoint(insertionPoint)

        # Process the Docx File
        self.process_doc(displayProgress)

    def SetTxtStyle(self, fontColor = None, fontBgColor = None, fontFace = None, fontSize = None,
                          fontBold = None, fontItalic = None, fontUnderline = None,
                          parAlign = None, parLeftIndent = None, parRightIndent = None,
                          parTabs = None, parLineSpacing = None, parSpacingBefore = None, parSpacingAfter = None):
        """ I find some of the RichTextCtrl method names to be misleading.  Some character styles are stacked in the RichTextCtrl,
            and they are removed in the reverse order from how they are added, regardless of the method called.

            For example, starting with plain text, BeginBold() makes it bold, and BeginItalic() makes it bold-italic. EndBold()
            should make it italic but instead makes it bold. EndItalic() takes us back to plain text by removing the bold.

            According to Julian, this functions "as expected" because of the way the RichTextCtrl is written.

            The SetTxtStyle() method handles overlapping styles in a way that avoids this problem.  """

        # If the font face (font name) is specified, set the font face
        if fontFace:
            self.txtAttr.SetFontFaceName(fontFace)
            self.font['fontfacename'] = fontFace

        # If the font size is specified, set the font size
        if fontSize:
            self.txtAttr.SetFontSize(fontSize)
            self.font['fontpointsize'] = fontSize

        # If a color is specified, set text color
        if fontColor:
            self.txtAttr.SetTextColour(fontColor)
        # If a background color is specified, set the background color
        if fontBgColor:
            self.txtAttr.SetBackgroundColour(fontBgColor)
        # If bold is specified, set or remove bold as requested
        if fontBold != None:
            if fontBold:
                self.txtAttr.SetFontWeight(wx.FONTWEIGHT_BOLD)
            else:
                self.txtAttr.SetFontWeight(wx.FONTWEIGHT_NORMAL)
        # If italics is specified, set or remove bold as requested
        if fontItalic != None:
            if fontItalic:
                self.txtAttr.SetFontStyle(wx.FONTSTYLE_ITALIC)
            else:
                self.txtAttr.SetFontStyle(wx.FONTSTYLE_NORMAL)
        # If underline is specified, set or remove bold as requested
        if fontUnderline != None:
            if fontUnderline:
                self.txtAttr.SetFontUnderlined(True)
            else:
                self.txtAttr.SetFontUnderlined(False)
        # If Paragraph Alignment is specified, set the alignment
        if parAlign != None:
            self.txtAttr.SetAlignment(parAlign)
        # If Left Indent is specified, set the left indent
        if parLeftIndent != None:
            # Left Indent can be an integer for left margin only, or a 2-element tuple for left indent and left subindent.
            if type(parLeftIndent) == int:
                self.txtAttr.SetLeftIndent(parLeftIndent)
            elif (type(parLeftIndent) == tuple) and (len(parLeftIndent) > 1):
                self.txtAttr.SetLeftIndent(parLeftIndent[0], parLeftIndent[1])
        # If Right Indent is specified, set the right indent
        if parRightIndent != None:
            self.txtAttr.SetRightIndent(parRightIndent)
        # If Tabs are specified, set the tabs
        if parTabs != None:
            self.txtAttr.SetTabs(parTabs)
        # If Line Spacing is specified, set Line Spacing
        if parLineSpacing != None:
            self.txtAttr.SetLineSpacing(parLineSpacing)
        # If Paragraph Spacing Before is set, set spacing before
        if parSpacingBefore != None:
            self.txtAttr.SetParagraphSpacingBefore(parSpacingBefore)
        # If Paragraph Spacing After is set, set spacing after
        if parSpacingAfter != None:
            self.txtAttr.SetParagraphSpacingAfter(parSpacingAfter)
        # Apply the modified font to the document
        self.txtCtrl.SetDefaultStyle(self.txtAttr)
      
    def process_doc(self, displayProgress=True):
        """ Process and parse a document in Word Docx Format """
        readOnly = not self.txtCtrl.IsEditable()
        if readOnly:
            self.txtCtrl.SetEditable(True)

        if DEBUG:
            print "PyDocxParser.process_doc():", len(self.document.paragraphs)
            startTime = time.time()

        if displayProgress:

            progressDlg = wx.ProgressDialog(_("Parsing Docx Format Data"),
                                            _("Importing Docx Format document"),
                                            len(self.document.paragraphs), None,
                                            wx.PD_APP_MODAL | wx.PD_AUTO_HIDE |
                                            wx.PD_ELAPSED_TIME | wx.PD_ESTIMATED_TIME | wx.PD_REMAINING_TIME )
            progressDlg.CentreOnScreen()
        else:
            progressDlg = None

        pCount = 0

        enum = docx.enum.text.WD_COLOR_INDEX
        # Highlight Colors defined in Microsoft Word
        colorConversions = {enum.AUTO          :  None,                     # Word:  Automatic
                            enum.BLACK         :  wx.Colour(0, 0, 0),       # Word:  Black
                            enum.BLUE          :  wx.Colour(0, 0, 255),     # Word:  Blue
                            enum.BRIGHT_GREEN  :  wx.Colour(0, 255, 0),     # Word:  Bright Green
                            enum.DARK_BLUE     :  wx.Colour(0, 0, 139),     # Word:  Dark Blue
                            enum.DARK_RED      :  wx.Colour(139, 0, 0),     # Word:  Dark Red
                            enum.DARK_YELLOW   :  wx.Colour(128, 128, 0),   # Word:  Dark Yellow
                            enum.GRAY_25       :  wx.Colour(211, 211, 211), # Word:  Gray 25%
                            enum.GRAY_50       :  wx.Colour(169, 169, 169), # Word:  Gray 50%
                            enum.GREEN         :  wx.Colour(0, 100, 0),     # Word:  Green
                            enum.PINK          :  wx.Colour(255, 0, 255),   # Word:  Pink
                            enum.RED           :  wx.Colour(255, 0, 0),     # Word:  Red
                            enum.TEAL          :  wx.Colour(0, 139, 139),   # Word:  Teal
                            enum.TURQUOISE     :  wx.Colour(0, 255, 255),   # Word:  Turquoise
                            enum.VIOLET        :  wx.Colour(128, 0, 128),   # Word:  Violet
                            enum.WHITE         :  wx.Colour(255, 255, 255), # Word:  White
                            enum.YELLOW        :  wx.Colour(255, 255, 0)}   # Word:  Yellow

        for p in self.document.paragraphs:

            pCount += 1
            
            if p.paragraph_format.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER:
                pParAlign = wx.TEXT_ALIGNMENT_CENTER
            elif p.paragraph_format.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT:
                pParAlign = wx.TEXT_ALIGNMENT_RIGHT
            else:
                pParAlign = wx.TEXT_ALIGNMENT_LEFT

            if p.paragraph_format.left_indent != None:
                pParLeftIndent = p.paragraph_format.left_indent / 3556.0
            else:
                pParLeftIndent = 0

            if p.paragraph_format.first_line_indent != None:
                pParFirstLineIndent = p.paragraph_format.first_line_indent / 3556.0
            else:
                pParFirstLineIndent = 0
                
            if p.paragraph_format.right_indent != None:
                pParRightIndent = p.paragraph_format.right_indent / 3556.0
            else:
                pParRightIndent = 0

            if p.paragraph_format.line_spacing != None:
                pLineSpacing = p.paragraph_format.line_spacing * 10
            else:
                pLineSpacing = 0

            if p.paragraph_format.space_before != None:
                pSpacingBefore = p.paragraph_format.space_before / 3556.0
            else:
                pSpacingBefore = 0

            if p.paragraph_format.space_after != None:
                pSpacingAfter = p.paragraph_format.space_after / 3556.0
            else:
                pSpacingAfter = 35  # Word defaults to 10pt, which is 35!

            pTabs = []
            if len(p.paragraph_format.tab_stops) > 1:
                for tab in p.paragraph_format.tab_stops:
                    # Tab stops in the RichTextCtrl are stored in 1/10ths of a millimeter
                    pTabs.append(int(tab.position / 3600.0))

            self.SetTxtStyle(parAlign = pParAlign,
                             parLeftIndent = (pParFirstLineIndent + pParLeftIndent, 0 - pParFirstLineIndent),
                             parRightIndent = pParRightIndent,
                             parLineSpacing = pLineSpacing,
                             parSpacingBefore = pSpacingBefore,
                             parSpacingAfter = pSpacingAfter,
                             parTabs = pTabs)
            
            for r in p.runs:

                if r.font.size != None:
                    rFontSize = r.font.size / 12700
                else:
                    # None indicates no change!  NO NONE MEANS USE WORD DEFAULT!!
                    rFontSize = 11  # Word defaults to 11 pt, I believe
                    
                if r.font.bold:
                    rFontBold = True
                else:
                    rFontBold = False
                if r.font.italic:
                    rFontItalic = True
                else:
                    rFontItalic = False
                if r.font.underline:
                    rFontUnderline = True
                else:
                    rFontUnderline = False

                if r.font.color.rgb != None:
                    txtColor = r.font.color.rgb.__str__()
                else:
                    txtColor = '000000'

                if r.font.highlight_color != None:

                    bgColor = r.font.highlight_color

                    if bgColor in colorConversions.keys():
                        bgColor = colorConversions[bgColor]
                    else:
                        # Fail-over background color is white
                        bgColor = wx.Colour(255, 255, 255)
                else:
                    # Fail-over background color is white
                    bgColor = wx.Colour(255, 255, 255)
                self.SetTxtStyle(fontFace = r.font.name, fontSize = rFontSize,
                                 fontBold = rFontBold, fontItalic = rFontItalic, fontUnderline = rFontUnderline,
                                 fontColor = "#%s" % txtColor,
                                 fontBgColor = bgColor)

                hasHighlight_color = hasattr(r.font, 'highlight_color')

                # Start Exception Handling
                try:
                    # If we have text ...
                    if r.text != '':
                        # ... then add that text to the wxRichTextCtrl.
                        self.txtCtrl.WriteText(r.text)
                    # If we have a picture ...
                    elif r.has_picture:
                        # This quick dictionary allows us to translate between *.docx image types and wxPython image types
                        imageTypes = { 'image/png'  : wx.BITMAP_TYPE_PNG,
                                       'image/jpeg' : wx.BITMAP_TYPE_JPEG,
                                       'image/gif'  : wx.BITMAP_TYPE_GIF }

                        # For each picture in the picture list (although I've never seen more than one!)
                        for picture in r.picture_lst:
                            # ... if we have an image of a known type (PNG and JPEG) ...
                            if picture.image_type in imageTypes.keys():
                                # ... get a StringIO stream to take the image data.
                                stream = cStringIO.StringIO(picture.image_data)
                                # Now convert that stream to an image
                                img = wx.ImageFromStream(stream, imageTypes[picture.image_type])
                                # If we were successful in creating a valid image ...
                                if img.IsOk():
                                    # ... add that image to the wxRichTextEdit control
                                    self.txtCtrl.WriteImage(img)
                            # If the image type is unknown, I can probably fix that!  Put a message in the text!!
                            else:
                                self.txtCtrl.WriteText(' (( Unknown Image Type: "%s" )) ' % picture.image_type)
                    # If we have neither text nor a picture ...
                    else:
                        # ... we could put a message in the text, or not.
                        if False:
                            self.txtCtrl.WriteText(' (( r.text is blank and r.has_picture is False )) ')
                # If we get a UnicodeDecodeError ...
                except UnicodeDecodeError:
                    # ... put a SPACE in the Transcript
                    self.txtCtrl.WriteText(' ')

                    # ... and put a note in the Error Log!
                    print "PyDocxParser.DocxTowxRichTextCtrlParser.process_txt():  UnicodeDecodeError:", len(r.text),
                    if len(r.text) == 1:
                        print ord(r.text)
                    else:
                        for x in r.text:
                            print ord(x),
                        print

            self.txtCtrl.WriteText('\n')

        # If we started out in Read Only mode ...
        if readOnly:
            # ... then return to Read Only mode
            self.txtCtrl.SetReadOnly(True)

        if progressDlg:
            progressDlg.Close()
            progressDlg.Destroy()

            wx.YieldIfNeeded()

        if DEBUG:
            print "Exiting PyDocxParser.DocxTowxRichTextCtrlParser.process_doc():", time.time() - startTime


# If we're running in stand-alone test mode
if __name__ == '__main__':
    # Create an xml.sax parser
    parser = xml.sax.make_parser()
    # Define our XML to DOCx Handler
    handler = XMLToDocxHandler()
    # Set the parser to use the handler
    parser.setContentHandler(handler)
    # Open a test XML file, 'test.xml', which should be created by saving XML from a wxRichTextCtrl, and parse it
    parser.parse("test.xml")
    # Save the resulting DOCx string to a file called 'text.rtf'
    handler.saveFile("test.docx")
